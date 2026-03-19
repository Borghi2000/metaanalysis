from docx import Document
import re

doc_path = r"e:\Pesquisa projeto\PDF DA PESQUISA\paper_fusao_v10_criticas.docx"
doc = Document(doc_path)

print("--- FIGURE SEARCH ---")
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if re.search(r"(Figura|Figure|Fig\.|SROC|Forest|Fagan)", text, re.I):
        # Print paragraph index, text, and 50 chars of context if it's a long paragraph
        print(f"[{i}]: {text[:200]}...")

print("\n--- SECTION HEADERS ---")
for i, para in enumerate(doc.paragraphs):
    if para.style.name.startswith('Heading') or (para.text.isupper() and len(para.text) < 50):
         print(f"[{i}] {para.style.name}: {para.text}")
