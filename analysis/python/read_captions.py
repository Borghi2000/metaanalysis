from docx import Document

doc_path = r"e:\Pesquisa projeto\PDF DA PESQUISA\paper_fusao_v10_criticas.docx"
doc = Document(doc_path)

print("--- FIGURE CAPTIONS ---")
for i, para in enumerate(doc.paragraphs):
    if "⬛  INSERIR FIGURA" in para.text:
        caption = ""
        # Look back up to 3 paragraphs for the caption starting with 'Figura'
        for j in range(1, 4):
            prev = doc.paragraphs[i-j].text.strip()
            if prev.startswith("Figura") or prev.startswith("Figure"):
                caption = prev
                break
        print(f"Placeholder na linha {i}: {para.text.strip()} => Caption: {caption}")
